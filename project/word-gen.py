from docx import Document
import json

# 创建一个新文档
doc = Document()

# 添加标题
doc.add_heading('一. 单选题', level=1)

# 添加段落内容
with open('./project/single_select_data.json', "r") as file:
    data = json.load(file)
    print(data)
    # doc.add_paragraph('这是一个由 Python 生成的 Word 文档。')

# 保存文档
# doc.save('sample.docx')